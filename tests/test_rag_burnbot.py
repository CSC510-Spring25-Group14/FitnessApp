import os
import pytest
from unittest.mock import patch, MagicMock
import rag_burnbot
from docx import Document
from rag_burnbot import (
    cleaned_text, extract_text_from_document, chunk_text, build_faiss_index, retrieve_context, gemini_response, bot_response
)
import numpy as np
import faiss

# Test 01: Extracting text from a document
def test_extract_text_from_document():
    path = "./sample_test.docx"
    doc = Document()
    doc.add_paragraph("First line")
    doc.add_paragraph("Second line")
    doc.save(path)

    text = extract_text_from_document(path)
    assert isinstance(text, list)
    assert text == ["First line", "Second line"]
    os.remove(path)
    
# Test 02: Checking if a document is empty
def test_extract_text_from_document_empty():
    path = "./empty.docx"
    doc = Document()
    doc.save(path)
    assert extract_text_from_document(path) == []
    os.remove(path)
    
# Test 03: Cleaning whitespace from a string
def test_cleaned_text_whitespaces():
    assert cleaned_text("    This is    a   test.    ") == "This is a test."

# Test 04: Cleaning text with new lines and tabs
def test_cleaned_text_newlines():
    assert cleaned_text("\n\tThis is a test.\n\nThis is another line.") == "This is a test. This is another line."

# Test 05: Extracting text from a document
@patch("rag_burnbot.Document")
def test_extract_text_from_document(mock_document):
    mock_document.return_value.paragraphs = [
        MagicMock(text="This is a test."),
        MagicMock(text="   "),
        MagicMock(text="This is a new line.")
    ]
    with patch("rag_burnbot.Document", return_value=mock_document.return_value):
        result = extract_text_from_document("fake_path.docx")
        assert result == ["This is a test.", "This is a new line."]
        
# Test 06: Testing chunk size is within limit
def test_chunk_text_size():
    paragraphs = ["a" * 200, "b" * 200, "c" * 200]
    chunks = chunk_text(paragraphs)
    assert all(len(chunk) <= 500 for chunk in chunks)
    
# Test 07: Testing chunk remainder
def test_chunk_text_remainder():
    paragraphs = ["short text"]
    chunks = chunk_text(paragraphs, chunk_size=100)
    assert chunks == ["short text"]
    
# Test 08: Testing Faiss vectors
def test_build_faiss_index():
    chunks = ["chunk1", "chunk2"]
    index, _ = build_faiss_index(chunks)
    assert index.ntotal == 2
    
# Test 09: Retrieving context
@patch("rag_burnbot.embedding_model.encode", return_value=np.random.rand(1, 384).astype("float32"))
def test_retrieve_context_returns_correct_chunks(mock_encode):
    chunks = ["This is about rice.", "This is about bananas."]
    index, chunk_store = build_faiss_index(chunks)
    context = retrieve_context("rice", index, chunk_store, k=1)
    assert isinstance(context, list)
    assert len(context) == 1
    
# Test 10: Testing the shape of the retrieved context
@patch("rag_burnbot.embedding_model.encode", return_value=np.random.rand(1, 384).astype("float32"))
def test_retrieve_context_shape(mock_encode):
    chunks = ["This is about rice.", "This is about bananas."]
    index, chunk_store = build_faiss_index(chunks)
    context = retrieve_context("some query", index, chunk_store, k=2)
    assert len(context) == 2
    
# Test 11: Testing invalid index of the context
def test_retrieve_context_invalid_index():
    with pytest.raises(AttributeError):
        retrieve_context("query", None, None)
        
# Test 12: Testing the generated content from model
@patch("rag_burnbot.model.generate_content")
def test_gemini_response_valid(mock_generate):
    mock_generate.return_value.text = "**Calories in rice are 200.**"
    response = gemini_response(["Rice is 200 calories."], "What are calories in rice?")
    assert "200" in response
    assert "*" not in response

# Test 13:    
@patch("rag_burnbot.model.generate_content", side_effect=Exception("API down"))
def test_gemini_response_failure(mock_generate):
    result = gemini_response(["context"], "question?")
    assert "Gemini API error" in result

# Test 14:   
def test_bot_response_menu_reset():
    chunks = ["Sample context"]
    index, chunk_store = build_faiss_index(chunks)
    for cmd in ["0", "start", "menu", "reset", "restart"]:
        assert "BurnBot" in bot_response(cmd, index, chunk_store)

# Test 15:         
@patch("rag_burnbot.retrieve_context", return_value=["Calories info"])
@patch("rag_burnbot.gemini_response", return_value="Calories in rice are 300.")
def test_bot_response_valid_query(mock_gemini, mock_context):
    chunks = ["Sample context"]
    index, chunk_store = build_faiss_index(chunks)
    response = bot_response("calories in rice?", index, chunk_store)
    assert "300" in response
    
# Test 16: 
def test_chunk_overlap_functionality():
    paragraphs = ["This is sentence one.", "This is sentence two."]
    chunks = chunk_text(paragraphs, chunk_size=40, chunk_overlap=10)
    assert len(chunks) >= 1
    
# Test 17: 
def test_chunk_boundary_content():
    paragraphs = ["Sentence one." * 30]  # Force large chunk
    chunks = chunk_text(paragraphs, chunk_size=200, chunk_overlap=50)
    assert len(chunks) >= 2
    
# Test 18:
@patch("rag_burnbot.retrieve_context", return_value=["Large context"])
@patch("rag_burnbot.gemini_response", return_value="Large input handled.")
def test_bot_response_large_query(mock_gemini, mock_context):
    chunks = ["Sample context"] * 5
    index, chunk_store = build_faiss_index(chunks)
    large_q = "What are the effects of carbs on workout?" * 10
    result = bot_response(large_q, index, chunk_store)
    assert "Large input handled." in result
    
# Test 19:
def test_faiss_index_wrong_shape():
    bad_vectors = np.random.rand(2, 100).astype("float32")
    index = faiss.IndexFlatL2(384)
    with pytest.raises(Exception):
        index.add(bad_vectors)
        
# Test 20:
@patch("rag_burnbot.embedding_model.encode", return_value=np.random.rand(1, 384).astype("float32"))
def test_retrieve_context_k_greater_than_chunks(mock_encode):
    chunks = ["A"] * 3
    index, chunk_store = build_faiss_index(chunks)
    context = retrieve_context("test", index, chunk_store, k=10)
    assert len(context) <= 10